"""
Document Parser — extracts raw text from uploaded PDF, DOCX, and TXT files.

Supported formats:
  - PDF  (.pdf)  via pypdf
  - DOCX (.docx) via python-docx
  - TXT  (.txt)  via plain UTF-8 read
"""

from __future__ import annotations

import logging
import os
from pathlib import Path
from typing import Optional

logger = logging.getLogger("org_knowledge.document_parser")


class DocumentParsingError(Exception):
    """Raised when a document cannot be parsed."""


def parse_document(file_path: str | Path) -> str:
    """
    Parse a document at the given path and return its text content.

    Args:
        file_path: Absolute or relative path to the document.

    Returns:
        Extracted plain text from the document.

    Raises:
        DocumentParsingError: If the file type is unsupported or parsing fails.
        FileNotFoundError: If the file does not exist.
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    ext = path.suffix.lower()

    if ext == ".pdf":
        return _parse_pdf(path)
    elif ext == ".docx":
        return _parse_docx(path)
    elif ext == ".txt":
        return _parse_txt(path)
    else:
        raise DocumentParsingError(
            f"Unsupported file format: '{ext}'. "
            f"Supported formats: .pdf, .docx, .txt"
        )


def parse_document_bytes(file_bytes: bytes, filename: str) -> str:
    """
    Parse a document from raw bytes.

    Args:
        file_bytes: The raw file content as bytes.
        filename: Original filename (used to determine the format).

    Returns:
        Extracted plain text.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _parse_pdf_bytes(file_bytes)
    elif ext == ".docx":
        return _parse_docx_bytes(file_bytes)
    elif ext == ".txt":
        return file_bytes.decode("utf-8", errors="replace")
    else:
        raise DocumentParsingError(
            f"Unsupported file format: '{ext}'. "
            f"Supported formats: .pdf, .docx, .txt"
        )


def _parse_pdf(path: Path) -> str:
    """Extract text from a PDF file using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise DocumentParsingError(
            "pypdf is not installed. Run: pip install pypdf"
        )

    try:
        reader = PdfReader(str(path))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)
    except Exception as exc:
        raise DocumentParsingError(f"Failed to parse PDF '{path.name}': {exc}")


def _parse_pdf_bytes(file_bytes: bytes) -> str:
    """Extract text from PDF bytes."""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise DocumentParsingError("pypdf is not installed. Run: pip install pypdf")

    try:
        import io
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
        return "\n\n".join(pages)
    except Exception as exc:
        raise DocumentParsingError(f"Failed to parse PDF from bytes: {exc}")


def _parse_docx(path: Path) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        raise DocumentParsingError(
            "python-docx is not installed. Run: pip install python-docx"
        )

    try:
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as exc:
        raise DocumentParsingError(f"Failed to parse DOCX '{path.name}': {exc}")


def _parse_docx_bytes(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    try:
        from docx import Document
    except ImportError:
        raise DocumentParsingError("python-docx is not installed. Run: pip install python-docx")

    try:
        import io
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return "\n".join(paragraphs)
    except Exception as exc:
        raise DocumentParsingError(f"Failed to parse DOCX from bytes: {exc}")


def _parse_txt(path: Path) -> str:
    """Read a plain text file with UTF-8 encoding."""
    try:
        return path.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        # Fallback to latin-1 if UTF-8 fails
        return path.read_text(encoding="latin-1")
    except Exception as exc:
        raise DocumentParsingError(f"Failed to read TXT '{path.name}': {exc}")

